#!/usr/bin/env python3
import argparse, io, json, hashlib, random, string, time
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Tuple

import boto3
from botocore.config import Config
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import LETTER
from tqdm import tqdm

# --------------------------
# Constants & taxonomies
# --------------------------

ICH_E3_SECTIONS = [
    "1. Title Page",
    "2. Synopsis",
    "3. Table of Contents",
    "4. Ethics Committee/IRB Information",
    "5. Investigators and Administrative Structure",
    "6. Introduction",
    "7. Study Objectives",
    "8. Investigational Plan",
    "9. Study Patients",
    "10. Efficacy Evaluation",
    "11. Safety Evaluation",
    "12. Discussion and Overall Conclusions",
    "13. Tables, Figures, and Graphs Referred to but Not Included",
    "14. Reference List",
    "15. Appendices",
    "16. Signatures",
]

# 100 representative TMF artifact codes/names (mix across sections 01–16)
TMF_ARTIFACTS: List[Tuple[str, str]] = [
    ("01.01.01", "Protocol"),
    ("01.01.02", "Protocol Amendment"),
    ("01.01.03", "Protocol Signature Page"),
    ("01.02.01", "Investigator’s Brochure"),
    ("01.03.01", "Sample CRF"),
    ("01.05.01", "Statistical Analysis Plan"),
    ("01.06.01", "Randomization Scheme"),
    ("01.07.01", "Blinding Procedures"),
    ("02.01.01", "IRB/IEC Approval"),
    ("02.02.01", "Informed Consent Approval"),
    ("03.01.01", "Regulatory Authority Approval"),
    ("03.02.01", "Clinical Trial Application"),
    ("03.03.01", "Safety Reporting Correspondence"),
    ("04.01.01", "Investigational Product Dossier"),
    ("04.02.01", "IMP Accountability Log"),
    ("04.03.01", "Temperature Excursion Report"),
    ("05.01.01", "Monitoring Plan"),
    ("05.01.07", "Monitoring Visit Report"),
    ("05.01.09", "Follow-up Letter to Site"),
    ("05.03.01", "Central Monitoring Reports"),
    ("06.01.01", "Informed Consent Form (Master)"),
    ("06.01.02", "Informed Consent Form (Site)"),
    ("06.02.01", "Consent Process Documentation"),
    ("07.01.01", "Investigator CV"),
    ("07.01.02", "Sub-Investigator CV"),
    ("07.02.01", "Financial Disclosure"),
    ("08.01.01", "Training Plan"),
    ("08.02.01", "Training Record"),
    ("08.03.01", "SOP Acknowledgement"),
    ("09.01.01", "Site Initiation Checklist"),
    ("09.02.01", "Site Activation Letter"),
    ("09.03.01", "Enrollment Log"),
    ("09.04.01", "Delegation of Authority Log"),
    ("09.05.01", "Screening Log"),
    ("09.06.01", "Signature Log"),
    ("10.01.01", "Laboratory Certification"),
    ("10.02.01", "Normal Ranges/Reference Values"),
    ("10.03.01", "Central Lab Manual"),
    ("11.01.01", "SAE/SUSAR Reports"),
    ("11.02.01", "Annual DSUR"),
    ("11.03.01", "Safety Management Plan"),
    ("12.01.01", "Data Management Plan"),
    ("12.02.01", "Edit Check Specification"),
    ("12.03.01", "Database Lock Memo"),
    ("12.04.01", "Data Transfer Agreement"),
    ("13.01.01", "Quality Management Plan"),
    ("13.02.01", "Audit Report"),
    ("13.03.01", "CAPA Plan"),
    ("13.04.01", "Risk Assessment"),
    ("14.01.01", "CSR (Final)"),
    ("14.01.02", "CSR Synopsis"),
    ("14.02.01", "Tables, Listings and Figures Index"),
    ("14.03.01", "Patient Listings"),
    ("14.03.02", "Efficacy Listings"),
    ("14.03.03", "Safety Listings"),
    ("15.01.01", "Anonymization Report"),
    ("15.02.01", "Redaction Justification"),
    ("15.03.01", "Document History Log"),
    ("16.01.01", "Final Study Report Link"),
    ("16.02.01", "Study Close-out Memo"),
    ("16.03.01", "Archive Certificate"),
    # repeat/extend variations to reach 100:
] + [
    (f"{i:02d}.{j:02d}.{k:02d}", f"Artifact {i:02d}-{j:02d}-{k:02d}")
    for i in range(1, 17)
    for j in range(1, 4)
    for k in range(1, 3)
][:100]  # ensure we have exactly 100 entries

LOREM = (
    "This is synthetic content generated for development and testing. "
    "It follows ICH E3 or TMF Reference Model structure but contains no PHI/PII. "
    "Use for internal testing of ingestion, parsing, search, and provenance."
)

# --------------------------
# Helpers
# --------------------------

def rand_paragraph(n_sentences=5) -> str:
    base = LOREM
    extra = " ".join(
        " ".join(random.choice(string.ascii_lowercase) for _ in range(8))
        for _ in range(n_sentences)
    )
    return f"{base} {extra}"

def make_docx_paragraph_doc(sections: List[str], header_title: str) -> bytes:
    doc = Document()
    doc.add_heading(header_title, level=1)
    for sec in sections:
        doc.add_heading(sec, level=2)
        for _ in range(3):
            doc.add_paragraph(rand_paragraph(6))
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def make_pdf_paragraph_doc(sections: List[str], header_title: str) -> bytes:
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=LETTER)
    width, height = LETTER
    y = height - 72
    def write_line(txt, size=12):
        nonlocal y
        if y < 72:
            c.showPage()
            y = height - 72
        c.setFont("Helvetica", size)
        c.drawString(72, y, txt[:1000])
        y -= 16

    write_line(header_title, size=14)
    for sec in sections:
        write_line("")
        write_line(sec, size=13)
        for _ in range(3):
            # wrap paragraph into lines:
            para = rand_paragraph(6)
            words = para.split()
            line = ""
            for w in words:
                if len(line) + len(w) + 1 > 100:
                    write_line(line)
                    line = w
                else:
                    line = (line + " " + w).strip()
            if line:
                write_line(line)
    c.showPage()
    c.save()
    return bio.getvalue()

def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

def s3_put_bytes(s3, bucket: str, key: str, data: bytes, content_type: str):
    s3.put_object(Bucket=bucket, Key=key, Body=data, ContentType=content_type)

def put_json(s3, bucket: str, key: str, obj: Dict):
    s3.put_object(Bucket=bucket, Key=key, Body=json.dumps(obj, indent=2).encode("utf-8"), ContentType="application/json")

# --------------------------
# Generators
# --------------------------

def generate_csr_docs(s3, bucket: str, prefix: str, count: int, versions: str, formats: List[str], region: str):
    base_date = datetime.utcnow() - timedelta(days=30)
    for i in tqdm(range(1, count + 1), desc="Generating CSRs"):
        study_id = f"STUDY-{i:04d}"
        version = versions
        header = f"Clinical Study Report — {study_id} — Version {version}"
        # files
        docx_b = make_docx_paragraph_doc(ICH_E3_SECTIONS, header) if "docx" in formats else None
        pdf_b = make_pdf_paragraph_doc(ICH_E3_SECTIONS, header) if "pdf" in formats else None

        # choose one or both uploads
        uploaded_paths = []
        pages_estimate = 10 * (1 if pdf_b else 1)  # rough estimate for synthetic

        if pdf_b:
            key = f"{prefix}csr/raw/{study_id}/{version}/CSR_{study_id}_{version}.pdf"
            s3_put_bytes(s3, bucket, key, pdf_b, "application/pdf")
            uploaded_paths.append(f"s3://{bucket}/{key}")
            pdf_sha = sha256_bytes(pdf_b)
        else:
            pdf_sha = None

        if docx_b:
            key = f"{prefix}csr/raw/{study_id}/{version}/CSR_{study_id}_{version}.docx"
            s3_put_bytes(s3, bucket, key, docx_b, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            uploaded_paths.append(f"s3://{bucket}/{key}")
            docx_sha = sha256_bytes(docx_b)
        else:
            docx_sha = None

        checksum = pdf_sha or docx_sha or ""
        created_at = (base_date + timedelta(days=i % 15)).isoformat(timespec="seconds") + "Z"
        metadata = {
            "document_id": f"csr:{study_id}:{version}",
            "doc_type": "CSR",
            "study_id": study_id,
            "doc_version": version,
            "source_paths": uploaded_paths,
            "mime_types": [ "application/pdf" if pdf_b else None, "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if docx_b else None ],
            "pages": pages_estimate,
            "checksum_sha256": checksum,
            "created_at": created_at,
            "effective_date": created_at[:10],
            "source_system": "S3",
            "labels": ["synthetic", "ich_e3_aligned"],
            "security_level": "internal",
            "ich_e3_compliance": True,
            "tmf_ref_model_version": None,
            "pii_detected": False,
            "pii_fields": []
        }
        key_md = f"{prefix}csr/processed/{study_id}/{version}/metadata/metadata.json"
        put_json(s3, bucket, key_md, metadata)

def generate_tmf_docs(s3, bucket: str, prefix: str, count: int, versions: str, formats: List[str], tmf_ref_version: str, region: str):
    base_date = datetime.utcnow() - timedelta(days=20)
    # ensure we have at least `count` artifacts; cycle if needed
    for i in tqdm(range(1, count + 1), desc="Generating TMFs"):
        study_id = f"STUDY-{i:04d}"
        version = versions
        code, name = TMF_ARTIFACTS[(i - 1) % len(TMF_ARTIFACTS)]
        header = f"TMF Artifact — {name} ({code}) — {study_id} — Version {version}"
        sections = [f"{name} — Section {k}" for k in range(1, 6)]

        docx_b = make_docx_paragraph_doc(sections, header) if "docx" in formats else None
        pdf_b = make_pdf_paragraph_doc(sections, header) if "pdf" in formats else None

        uploaded_paths = []
        if pdf_b:
            key = f"{prefix}tmf/raw/{study_id}/{code}/{version}/{name.replace(' ', '')}_{study_id}_{version}.pdf"
            s3_put_bytes(s3, bucket, key, pdf_b, "application/pdf")
            uploaded_paths.append(f"s3://{bucket}/{key}")
            pdf_sha = sha256_bytes(pdf_b)
        else:
            pdf_sha = None

        if docx_b:
            key = f"{prefix}tmf/raw/{study_id}/{code}/{version}/{name.replace(' ', '')}_{study_id}_{version}.docx"
            s3_put_bytes(s3, bucket, key, docx_b, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            uploaded_paths.append(f"s3://{bucket}/{key}")
            docx_sha = sha256_bytes(docx_b)
        else:
            docx_sha = None

        checksum = pdf_sha or docx_sha or ""
        created_at = (base_date + timedelta(days=i % 10)).isoformat(timespec="seconds") + "Z"
        metadata = {
            "document_id": f"tmf:{study_id}:{code}:{version}",
            "doc_type": "TMF",
            "study_id": study_id,
            "artifact_code": code,
            "artifact_name": name,
            "doc_version": version,
            "source_paths": uploaded_paths,
            "mime_types": [ "application/pdf" if pdf_b else None, "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if docx_b else None ],
            "tmf_ref_model_version": str(tmf_ref_version),
            "pages": 6,  # synthetic estimate
            "checksum_sha256": checksum,
            "created_at": created_at,
            "effective_date": created_at[:10],
            "source_system": "S3",
            "labels": ["synthetic", "tmf_ref_model_aligned"],
            "security_level": "internal",
            "ich_e3_compliance": False,
            "pii_detected": False,
            "pii_fields": []
        }
        key_md = f"{prefix}tmf/processed/{study_id}/{code}/{version}/metadata/metadata.json"
        put_json(s3, bucket, key_md, metadata)

# --------------------------
# CLI
# --------------------------

def main():
    ap = argparse.ArgumentParser(description="Generate synthetic CSR/TMF docs + metadata and upload to S3.")
    ap.add_argument("--bucket", required=True, help="Target S3 bucket")
    ap.add_argument("--prefix", default="", help="Optional key prefix (e.g., 'dev/')")
    ap.add_argument("--region", default="us-east-1")
    ap.add_argument("--csr-count", type=int, default=100)
    ap.add_argument("--tmf-count", type=int, default=100)
    ap.add_argument("--formats", nargs="+", default=["pdf"], choices=["pdf", "docx"])
    ap.add_argument("--version", default="v1")
    ap.add_argument("--tmf-ref-version", default="3.3")
    ap.add_argument("--profile", default=None, help="AWS profile (optional)")
    args = ap.parse_args()

    session_kwargs = {}
    if args.profile:
        session_kwargs["profile_name"] = args.profile
    session = boto3.Session(**session_kwargs)
    s3 = session.client("s3", region_name=args.region, config=Config(s3={"addressing_style": "virtual"}))

    # normalize prefix
    prefix = args.prefix.strip("/")
    if prefix:
        prefix = prefix + "/"

    generate_csr_docs(
        s3=s3,
        bucket=args.bucket,
        prefix=prefix,
        count=args.csr_count,
        versions=args.version,
        formats=args.formats,
        region=args.region,
    )
    generate_tmf_docs(
        s3=s3,
        bucket=args.bucket,
        prefix=prefix,
        count=args.tmf_count,
        versions=args.version,
        formats=args.formats,
        tmf_ref_version=args.tmf_ref_version,
        region=args.region,
    )
    print("✅ Done. Synthetic CSR & TMF samples uploaded with metadata.")

if __name__ == "__main__":
    main()
